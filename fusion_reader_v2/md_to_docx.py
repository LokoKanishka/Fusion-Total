import sys
import argparse
import re
import unicodedata
from pathlib import Path
from collections import Counter

SPANISH_FUNCTION_WORDS_V4 = {
    "a", "al", "ante", "bajo", "con", "contra", "de", "del", "desde", "durante",
    "e", "el", "ella", "ellas", "ellos", "en", "entre", "ha", "he", "hasta",
    "la", "las", "le", "lo", "los", "me", "mi", "mis", "ni", "no", "o", "para",
    "por", "que", "se", "si", "sin", "sobre", "son", "su", "sus", "te", "tu",
    "un", "una", "unas", "uno", "unos", "y", "ya",
}

PROTECTED_TERMS_V4 = {
    "Ars", "Magica", "Ars Magica", "Bonisagus", "Bjornaer", "Quaesitor",
    "Quaesitores", "Intellego", "Creo", "Muto", "Perdo", "Rego", "Corpus",
    "Mentem", "Animal", "Aquam", "Auram", "Ignem", "Terram", "Vim", "Jerbiton",
    "Flambeau", "Tremere", "Tytalus", "Verditius", "Criamon", "Merinita",
    "Miscellanea", "Ex Miscellanea", "Hermes", "Hermética", "Hermetica",
    "Voluntas", "Blackthorn", "Semitae", "Schola", "Pythagoranis", "Ungulus",
    "Cistercienses", "Rievaulx", "Stonehenge", "Guernicus", "Mercere",
    "Trianoma", "Antoninus", "Jocelin", "Fulk", "Wizards",
}

# Backward-compatible name used by older tests.
PROTECTED_TERMS = {term for term in PROTECTED_TERMS_V4 if " " not in term}

VALID_LONG_SPANISH_WORDS_V4 = {
    "administración", "administracion", "acontecimiento", "acontecimientos",
    "adelante", "adiestramiento", "afortunadamente", "agradecimientos",
    "alternativamente", "aparentemente",
    "apropiadamente", "automáticamente", "automaticamente", "característica",
    "características", "caracteristica", "caracteristicas", "completamente",
    "comportamiento", "considerablemente", "constantemente", "descubrimiento",
    "descubrimientos", "desgraciadamente", "especialidades", "especialmente",
    "específicamente", "especificamente", "experimentación", "experimentacion",
    "delante", "deliberadamente", "emocionalmente", "entretenimiento",
    "experimentador", "extremadamente", "excepcionalmente", "frecuentemente",
    "increíblemente", "increiblemente", "independientemente",
    "individualmente", "inevitablemente", "inherentemente", "inmediatamente", "instantáneamente",
    "instantaneamente", "intuitivamente", "laboratorio", "ligeramente",
    "naturalmente", "necesariamente", "ocasionalmente", "particularmente",
    "penalizaciones", "permanentemente", "personalidades", "principalmente",
    "posteriormente", "progresivamente", "reconocimiento", "representación", "representacion",
    "representaciones", "representantes", "responsabilidad", "restablecimiento",
    "respectivamente", "simultáneamente", "simultaneamente", "suficientemente",
    "verdaderamente", "voluntariamente", "envejecimiento", "encantamientos",
    "emplazamientos", "complementarios", "requerimientos", "pronunciamientos",
    "sobrenaturales", "incandescentes", "inconsistencias", "relacionándose", "relacionandose",
    "modelos", "adelard",
}

OCR_CANONICAL_WORDS_V4 = {
    "accion": "acción", "ademas": "además", "algun": "algún", "ano": "año",
    "anos": "años", "apendice": "apéndice", "asi": "así", "automaticamente": "automáticamente",
    "capitulo": "capítulo", "caracteristica": "característica",
    "caracteristicas": "características", "cogia": "cogía", "cogio": "cogió",
    "comence": "comencé", "companerismo": "compañerismo", "companero": "compañero",
    "companeros": "compañeros", "coordinacion": "coordinación", "dano": "daño",
    "despues": "después", "disenado": "diseñado", "diseniado": "diseñado",
    "disenio": "diseño", "edicion": "edición", "energia": "energía",
    "habia": "había", "habian": "habían", "habria": "habría", "hechizo": "hechizo",
    "indice": "índice", "introduccion": "introducción", "latín": "latín",
    "magica": "mágica", "magicas": "mágicas", "maquetacion": "maquetación",
    "mecanica": "mecánica", "mitica": "mítica", "mistico": "místico",
    "misticos": "místicos", "murio": "murió", "narracion": "narración",
    "numero": "número", "observé": "observé", "ocurrio": "ocurrió",
    "puntuacion": "puntuación", "senor": "señor", "senordemil": "señor de mil",
    "tambien": "también", "tecnica": "técnica", "tecnicas": "técnicas",
    "traduccion": "traducción", "volvi": "volví",
}

EXACT_GLUED_REPAIRS_V4 = {
    "Diariode": "Diario de",
    "Desarrollodela": "Desarrollo de la",
    "Edicionen": "Edición en",
    "delaspruebas": "de las pruebas",
    "Coordinaciondela": "Coordinación de la",
    "Majestuosoqueha": "Majestuoso que ha",
    "Senordemil": "Señor de mil",
    "cientoochentayseis": "ciento ochenta y seis",
    "yfue": "y fue",
    "yfuelavoluntad": "y fue la voluntad",
    "mediocomonovicioalmonasterio": "medio como novicio al monasterio",
    "lugardondehevivido": "lugar donde he vivido",
    "desdeentonces": "desde entonces",
    "Tomarémisvotoscomomonjelaproxima": "Tomaré mis votos como monje la próxima",
    "siemprehasostenidofirmementequelosmonjesdebenquedarse": "siempre ha sostenido firmemente que los monjes deben quedarse",
    "ymedijo": "y me dijo",
    "vestidoconelsimplehabitoblanco": "vestido con el simple hábito blanco",
    "Laspalabrasrompieroncualquierhechizoenelqueme": "Las palabras rompieron cualquier hechizo en el que me",
    "encontrarasumidoymevolviparahuir": "encontrara sumido y me volví para huir",
    "devueltaal": "de vuelta al",
    "campamentohabian": "campamento habían",
    "cuidadosamenteparaver": "cuidadosamente para ver",
    "sivolvia": "si volvía",
    "ynoesmuchomayorqueyo": "y no es mucho mayor que yo",
    "noesmuchomayorqueyo": "no es mucho mayor que yo",
    "Nuestroviaje": "Nuestro viaje",
    "representarenvuestrosrelatos": "representar en vuestros relatos",
    "suspoderesmisticosnohacensino": "sus poderes místicos no hacen sino",
    "incapacidadpara": "incapacidad para",
    "depersonas": "de personas",
    "oanimalesmundanos": "o animales mundanos",
    "oanimalesmundanossufrenunapenalizacion": "o animales mundanos sufren una penalización",
    "lugaresllamados": "lugares llamados",
    "clispersas": "dispersas",
    "tocda": "toda",
    "considerablegradoevariedad": "considerable grado y variedad",
    "unsimbolodelapertenencia": "un símbolo de la pertenencia",
    "Magoala": "Mago a la",
    "hechizosusanco": "hechizos usando",
    "Artesmagicas": "Artes mágicas",
    "unverho": "un verbo",
    "accion.y": "acción y",
    "Formulaicossonlos": "Formulaicos son los",
    "quehan": "que han",
    "efectospuedensermuyduradleros": "efectos pueden ser muy duraderos",
    "sonunsubconjunto": "son un subconjunto",
    "normalmentenumero": "normalmente número",
    "Nivelesde": "Niveles de",
    "Vidamuestran": "Vida muestran",
    "loheridoqueesta": "lo herido que está",
    "secontrolandemanerasimilara": "se controlan de manera similar a",
    "lasreglas": "las reglas",
    "quemanera": "qué manera",
    "elresultadoes": "el resultado es",
    "puedesvolver": "puedes volver",
    "vuelveatirar": "vuelve a tirar",
    "ycuadruplica": "y cuadruplica",
    "elresultado": "el resultado",
    "untercer": "un tercer",
    "multiplicaelresultadopor": "multiplica el resultado por",
    "habriatenidoquetirarundadodepifia": "habría tenido que tirar un dado de pifia",
    "incrementalaseveridad": "incrementa la severidad",
    "detufracaso": "de tu fracaso",
    "resultadodiferente": "resultado diferente",
    "Queelcontextoyelsentidocomunseantuguia": "Que el contexto y el sentido común sean tu guía",
    "CuartaEdicion": "Cuarta Edición",
    "ediciony": "edición y",
    "Wizardsofthe": "Wizards of the",
    "todoslos": "todos los",
    "todaslas": "todas las",
    "enel": "en el",
    "enla": "en la",
    "enlas": "en las",
    "enlos": "en los",
    "dela": "de la",
    "delas": "de las",
    "delos": "de los",
    "conel": "con el",
    "conla": "con la",
    "conlos": "con los",
    "porla": "por la",
    "porlo": "por lo",
    "paraque": "para que",
    "queha": "que ha",
    "queel": "que el",
    "queme": "que me",
    "quelos": "que los",
    "quelas": "que las",
    "conlas": "con las",
    "delaboratorio": "de laboratorio",
    "delarea": "del área",
    "delasumma": "de la summa",
    "DELARMA": "DEL ARMA",
    "mimaestra": "mi maestra",
    "mipadre": "mi padre",
    "mivida": "mi vida",
    "todamivida": "toda mi vida",
    "observadotodamivida": "observado toda mi vida",
    "estemundo": "este mundo",
    "estelibro": "este libro",
    "Simbolodelacasa": "Símbolo de la casa",
    "Elsimbolodelacasa": "El símbolo de la casa",
    "resistenciamagica": "resistencia mágica",
    "Criadoenel": "Criado en el",
    "delavictima": "de la víctima",
    "Igualqueel": "Igual que el",
    "unapuntuacionminima": "una puntuación mínima",
    "puntosdeexperiencia": "puntos de experiencia",
    "Eresespecialmentesensible": "Eres especialmente sensible",
    "Laboratoriosuperael": "Laboratorio supera el",
    "Familiarobtieneun": "Familiar obtiene un",
    "elobjetivodeestehechizo": "el objetivo de este hechizo",
    "delospersonajes": "de los personajes",
    "redondeadohacia": "redondeado hacia",
    "Algunoshechizosde": "Algunos hechizos de",
    "inmediatamentey": "inmediatamente y",
    "lanzamientorapido": "lanzamiento rápido",
    "todosloshechizos": "todos los hechizos",
}

def remove_image_placeholders(markdown: str) -> str:
    """Remove all image-related placeholders and markers."""
    markdown = re.sub(r'!\[.*?\]\(.*?\)', '', markdown)
    markdown = re.sub(r'<img.*?>', '', markdown, flags=re.IGNORECASE | re.DOTALL)
    markdown = re.sub(r'<!--.*?image.*?-->', '', markdown, flags=re.IGNORECASE)
    markdown = re.sub(r'data:image\/[a-zA-Z]*;base64,[a-zA-Z0-9+/=]*', '', markdown)
    return markdown


def _fold_spanish_v4(text: str) -> str:
    normalized = unicodedata.normalize("NFD", text.casefold())
    return "".join(ch for ch in normalized if unicodedata.category(ch) != "Mn")


PROTECTED_TERMS_FOLDED_V4 = {_fold_spanish_v4(term) for term in PROTECTED_TERMS_V4 if " " not in term}
VALID_LONG_FOLDED_V4 = {_fold_spanish_v4(word) for word in VALID_LONG_SPANISH_WORDS_V4}
EXACT_GLUED_REPAIRS_FOLDED_V4 = {_fold_spanish_v4(key): value for key, value in EXACT_GLUED_REPAIRS_V4.items()}


def _preserve_initial_case_v4(original: str, replacement: str) -> str:
    if not original:
        return replacement
    words = replacement.split()
    if not words:
        return replacement
    if original.isupper() and len(original) <= 5:
        words[0] = words[0].upper()
    elif original[0].isupper():
        words[0] = words[0][0].upper() + words[0][1:]
    return " ".join(words)


def _canonical_word_v4(word: str) -> str:
    return OCR_CANONICAL_WORDS_V4.get(_fold_spanish_v4(word), word)


def build_spanish_wordlist_v4() -> set[str]:
    """Build a local Spanish + Ars Magica wordlist for conservative splitting."""
    raw = """
    a accion acciones actualmente adelante adicional adicionalmente adivino alba abad abandonare abandonar
    abandonaran acorde acampamos acerca adiestramiento afin afuera afueras agua ahora aire ala al alba
    alabad alarde alianza alianzas aliados algun alguna algunas alguno algunos alma almonasterio alta
    ambiente animales ano anos ante antes antiguo antoninus apariencia apendice aprendido aprender aprencdiclos
    aquel aquella aquello aqui arquitectura arma artes asi asignado asunto ataque aunque autor aventura
    bajo basicas bendiga bendito bien blanco bonisagus bosque britania buena buenas caballo caballeros
    cada cambiar cambiado camino campamento capitulo capacidades cara carta casa causar centro cerca
    cistercienses ciudad clase claramente cogio comienza comenzo como comun comunicarse con conoce
    conocimiento confianza conjuro consortes construir contacto contra control controlar coordinacion corpus
    creo criamon cual cualquier cuando cuarta cuenta cuidad cuidadosamente dano dados dado de debe deben
    deberia decir decision defensa dejar dejara dejo del dentro derecho desde despues destino detectadas
    defectos dia diario dice diciendo dichoso dificultad dios direccion diseno disenio dispersas divina
    docena donde don doncellas dos durante duraderos ec efecto efectos edad edicion el ella ellos en
    encantamiento encantamientos encontrar encontrara energia enemigo enel entonces entre era eres especial
    espiritu esta estaba estan estar este esto estos europa ex experiencia externo familiar fatiga fe
    fenix ferico feerico firmemente forma formas formulaico formulaicos forima fracaso fulk general grado
    grogs guia ha habia habian habitoblanco habito hablar haces hacen hacer hacia han hasta he hechizo
    hechizos hermosa hermes hermetica hermetico herido historia hogar hombre hombres hora huir ignem
    iglesia imagen inmediatamente incapacidad incremento incrementa indicacion indice indivicduos informacion
    ingles inicio intellego invierno jerbiton jocelin juego jugadores junto juntos la laboratorio lanzador
    lanzar las latin lector lee libro llamados lugar magia magica magicas mago magos majestuoso manera
    manipulacion maquetacion maria mas mayor me mecanica medio mentem merinita mi mientras mil mis
    miscellanea mitica monasterio monje monjes momento mundo mundana mundanos muto narracion narrador
    nada nacimiento naturaleza necesitaras ni nivel niveles no nombre normales normalmente norte nos
    nosotros novicio nuestro nueva numero nunca objetivo observado observar ocaso ochenta ocurrio orden
    padre pagina palabras para parte partida partir pascua penalizacion pertenencia persona personaje
    personajes personas pifia pifias pluma poder poderes podria por porque posesiones precioso prueba
    pruebas puede pueden puesto puntuacion que quedara quedarse quemazon querer rancio realidad realizo
    recuerdo recuperaron reglas rego relacion relatos repetir representar resultado ritual rompieron
    saga sala salir sancta se seis semitae senor sentido ser sera sesion si siempre simple sino sistema
    sobre sobrenatural sociedad sostenido su sufrido sufrir sumido sus tabla tambien tecnicas tecnica
    tener tercero tierra tirar tocda toda todas todo todos tomar tomare traduccion tratar tremere tres
    trianoma tribunal tu tytalus un una usando usar ve ver verbo verdad verditius vestido vez viaje
    vida vigile vivido vivir volvi volver voluntad voluntas votar vuelve y ya
    diario antoninus jerbiton desarrollo edición coordinacion coordinación pruebas juego señor ciento
    ochenta seis monasterio novicio lugar vivido entonces sostenido firmemente monjes quedarse palabras
    rompieron cualquier hechizo encontrar sumido volví huir campamento habían vestido simple hábito
    blanco cistercienses alianza magos personaje personajes virtud defecto hermética laboratorio hechizos
    narración saga bestiario glosario latín diario desarrollo edición coordinación mágico mágica mágicas
    místico místicos número habría años año después algún también acción está había habían volvió
    """
    words = set(SPANISH_FUNCTION_WORDS_V4)
    words.update(_fold_spanish_v4(w) for w in raw.split())
    words.update(_fold_spanish_v4(w) for w in VALID_LONG_SPANISH_WORDS_V4)
    for replacement in EXACT_GLUED_REPAIRS_V4.values():
        words.update(_fold_spanish_v4(w) for w in re.findall(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+", replacement))
    words.update(_fold_spanish_v4(term) for term in PROTECTED_TERMS_V4 if " " not in term)
    for dic_path in (
        Path("/usr/share/hunspell/es.dic"),
        Path("/usr/share/calibre/dictionaries/es-ES/es-ES.dic"),
    ):
        if not dic_path.exists():
            continue
        try:
            for line in dic_path.read_text(encoding="utf-8", errors="ignore").splitlines():
                entry = line.strip().split("/", 1)[0]
                if entry and entry.isalpha() and 2 <= len(entry) <= 24:
                    words.add(_fold_spanish_v4(entry))
        except OSError:
            continue
    return {w for w in words if w}


SPANISH_WORDLIST_V4 = build_spanish_wordlist_v4()


def is_protected_term_v4(token: str) -> bool:
    folded = _fold_spanish_v4(token.strip())
    return folded in PROTECTED_TERMS_FOLDED_V4


def normalize_spanish_ocr_v4(text: str) -> tuple[str, dict]:
    """Conservative Spanish OCR accent repairs; no generative rewriting."""
    replacements = {
        r"\bHrs Magica\b": "Ars Magica",
        r"\bHRs Magica\b": "Ars Magica",
        r"\bARs magica\b": "Ars Magica",
        r"\bARs Magica\b": "Ars Magica",
        r"\braylagica\b": "Ars Magica",
        r"\bEdicion\b": "Edición",
        r"\bedicion\b": "edición",
        r"\bSenor\b": "Señor",
        r"\bsenor\b": "señor",
        r"\banio\b": "año",
        r"\bCoordinacion\b": "Coordinación",
        r"\bcoordinacion\b": "coordinación",
        r"\bMaquetacion\b": "Maquetación",
        r"\bmaquetacion\b": "maquetación",
        r"\bTraduccion\b": "Traducción",
        r"\btraduccion\b": "traducción",
        r"\bDireccion\b": "Dirección",
        r"\bdireccion\b": "dirección",
        r"\b(?:Introduccion|lntroduccion)\b": "Introducción",
        r"\bintroduccion\b": "introducción",
        r"\bCapitulo\b": "Capítulo",
        r"\bcapitulo\b": "capítulo",
        r"\bApendice\b": "Apéndice",
        r"\bapendice\b": "apéndice",
        r"\bIndice\b": "Índice",
        r"\bindice\b": "índice",
        r"\bMitica\b": "Mítica",
        r"\bmitica\b": "mítica",
        r"\bmecanica\b": "mecánica",
        r"\bMecanica\b": "Mecánica",
        r"\bpuntuacion\b": "puntuación",
        r"\bPuntuacion\b": "Puntuación",
        r"\bnarracion\b": "narración",
        r"\bNarracion\b": "Narración",
        r"\bcaracteristica\b": "característica",
        r"\bcaracteristicas\b": "características",
        r"\bcompanero\b": "compañero",
        r"\bcompaneros\b": "compañeros",
        r"\bdano\b": "daño",
        r"\bDano\b": "Daño",
        r"\bdiseniado\b": "diseñado",
        r"\bDiseniado\b": "Diseñado",
        r"\bdisenio\b": "diseño",
        r"\bDisenio\b": "Diseño",
        r"\bDetectos\b": "Defectos",
        r"\bHechbizos\b": "Hechizos",
        r"\bCuropa Mitica\b": "Europa Mítica",
        r"\bLatin\b": "Latín",
    }
    metrics = {"accent_fixes": 0}
    for pattern, replacement in replacements.items():
        text, count = re.subn(pattern, replacement, text)
        metrics["accent_fixes"] += count

    def magica_repl(match):
        start = max(0, match.start() - 5)
        if text[start:match.end()].lower().endswith("ars magica"):
            return match.group(0)
        return _preserve_initial_case_v4(match.group(0), "mágica")

    text, count = re.subn(r"\bmagica\b", magica_repl, text, flags=re.IGNORECASE)
    metrics["accent_fixes"] += count

    def temporal_ano(match):
        return _preserve_initial_case_v4(match.group(0), "año")

    text, count = re.subn(r"\bano\b(?=\s+(?:de|en|\d|mi|su|nuestro|octavo|primer|segundo|tercer))", temporal_ano, text, flags=re.IGNORECASE)
    metrics["accent_fixes"] += count
    return text, metrics


def normalize_common_ocr_errors(text: str) -> str:
    """Apply conservative corrections to common OCR mistakes."""
    text, _metrics = normalize_spanish_ocr_v4(text)
    return text

def _is_valid_long_word_v4(token: str) -> bool:
    return _fold_spanish_v4(token) in VALID_LONG_FOLDED_V4


def _is_suspicious_glued_token_v4(token: str) -> bool:
    folded = _fold_spanish_v4(token)
    if len(folded) < 5 or is_protected_term_v4(token) or _is_valid_long_word_v4(token):
        return False
    if folded in EXACT_GLUED_REPAIRS_FOLDED_V4:
        return True
    if re.search(r"[a-záéíóúüñ][A-ZÁÉÍÓÚÜÑ]", token):
        return True
    strong_patterns = (
        "dela", "delas", "delos", "enel", "enla", "enlas", "enlos", "queha",
        "queel", "queme", "quelos", "quelas", "conel", "conla", "conlos",
        "paraque", "porel", "porla", "desdeentonces", "demil", "ymedijo",
        "almonasterio", "dondehe", "comonovicio", "lasreglas", "resultadoes",
        "vuelvea", "sentidocomun", "artesmagicas", "poderesmisticos",
    )
    if any(pattern in folded for pattern in strong_patterns):
        return True
    return False


def detect_suspicious_glued_tokens(text: str) -> dict:
    tokens = re.findall(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+", text)
    suspicious: list[str] = []
    contexts: list[dict] = []
    protected_seen = sorted(
        {tok for tok in tokens if is_protected_term_v4(tok)},
        key=str.casefold,
    )
    for match in re.finditer(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+", text):
        token = match.group(0)
        if _is_suspicious_glued_token_v4(token):
            suspicious.append(token)
            if len(contexts) < 100:
                start = max(0, match.start() - 80)
                end = min(len(text), match.end() + 80)
                contexts.append({"token": token, "context": text[start:end].replace("\n", " ")})
    return {
        "total_tokens": len(tokens),
        "suspicious_count": len(suspicious),
        "suspicious_ratio": len(suspicious) / max(1, len(tokens)),
        "suspicious_tokens_top": Counter(suspicious).most_common(300),
        "suspicious_contexts": contexts,
        "protected_terms_seen": protected_seen,
    }


def _candidate_word_v4(part: str) -> str:
    folded = _fold_spanish_v4(part)
    return OCR_CANONICAL_WORDS_V4.get(folded, part)


def segment_glued_token_v4(token: str, wordlist: set[str] | None = None) -> tuple[str, bool, float]:
    """Split one suspicious Spanish OCR token with DP; returns text, changed, confidence."""
    if not token or is_protected_term_v4(token) or _is_valid_long_word_v4(token):
        return token, False, 0.0
    wordlist = wordlist or SPANISH_WORDLIST_V4
    folded = _fold_spanish_v4(token)
    if folded in EXACT_GLUED_REPAIRS_FOLDED_V4:
        return _preserve_initial_case_v4(token, EXACT_GLUED_REPAIRS_FOLDED_V4[folded]), True, 1.0
    if not _is_suspicious_glued_token_v4(token):
        return token, False, 0.0

    n = len(folded)
    best: list[tuple[float, list[str]] | None] = [None] * (n + 1)
    best[0] = (0.0, [])
    for i in range(n):
        if best[i] is None:
            continue
        score, parts = best[i]
        for j in range(i + 1, min(n, i + 18) + 1):
            piece = folded[i:j]
            if piece not in wordlist:
                continue
            original_piece = token[i:j]
            length = j - i
            if length <= 2 and piece not in SPANISH_FUNCTION_WORDS_V4:
                continue
            part_score = length * 2.0
            if piece in SPANISH_FUNCTION_WORDS_V4:
                part_score -= 1.2 if length <= 2 else 0.4
            if length >= 5:
                part_score += 2.0
            new_parts = parts + [_candidate_word_v4(original_piece)]
            cut_penalty = max(0, len(new_parts) - 1) * 2.7
            one_two = sum(1 for p in new_parts if len(_fold_spanish_v4(p)) <= 2)
            new_score = score + part_score - cut_penalty - (one_two * 0.7)
            if best[j] is None or new_score > best[j][0]:
                best[j] = (new_score, new_parts)
    if best[n] is None:
        return token, False, 0.0
    score, parts = best[n]
    if len(parts) < 2:
        return token, False, 0.0
    short_parts = sum(1 for p in parts if len(_fold_spanish_v4(p)) <= 2)
    confidence = min(0.99, max(0.0, score / max(12.0, len(folded) * 2.0)))
    if confidence < 0.42 or short_parts > max(3, len(parts) // 2 + 1):
        return token, False, confidence
    segmented = " ".join(parts)
    return _preserve_initial_case_v4(token, segmented), True, confidence


CONNECTOR_SPLITS_V4 = (
    ("desdeentonces", "desde entonces"),
    ("paraque", "para que"),
    ("delas", "de las"),
    ("delos", "de los"),
    ("dela", "de la"),
    ("enlas", "en las"),
    ("enlos", "en los"),
    ("enla", "en la"),
    ("enel", "en el"),
    ("conlas", "con las"),
    ("conlos", "con los"),
    ("conla", "con la"),
    ("conel", "con el"),
    ("porlas", "por las"),
    ("porlos", "por los"),
    ("porla", "por la"),
    ("porel", "por el"),
    ("quelas", "que las"),
    ("quelos", "que los"),
    ("queel", "que el"),
    ("queha", "que ha"),
    ("queme", "que me"),
)


def _split_connector_span_v4(token: str) -> tuple[str, bool]:
    folded = _fold_spanish_v4(token)
    if folded in {"modelos", "adelard"}:
        return token, False
    for pattern, replacement in CONNECTOR_SPLITS_V4:
        index = folded.find(pattern)
        if index < 0:
            continue
        before = token[:index]
        after = token[index + len(pattern):]
        if not before and not after:
            return _preserve_initial_case_v4(token, replacement), True
        if before and len(before) < 3:
            continue
        if after and len(after) < 2:
            continue
        pieces = []
        if before:
            pieces.append(before)
        pieces.append(replacement)
        if after:
            pieces.append(after)
        return _preserve_initial_case_v4(token, " ".join(pieces)), True
    return token, False


def repair_glued_words_v4(text: str) -> tuple[str, dict]:
    before = detect_suspicious_glued_tokens(text)
    text, ocr_metrics = normalize_spanish_ocr_v4(text)
    metrics = {
        "exact_fixes": 0,
        "dynamic_segmentations": 0,
        "accent_fixes": ocr_metrics.get("accent_fixes", 0),
        "protected_terms_skipped": 0,
        "low_confidence_skipped": 0,
        "suspicious_before": before["suspicious_count"],
        "suspicious_after": 0,
        "suspicious_reduction_percent": 0.0,
    }

    def repair_match(match):
        token = match.group(0)
        if is_protected_term_v4(token):
            metrics["protected_terms_skipped"] += 1
            return token
        folded = _fold_spanish_v4(token)
        if folded in EXACT_GLUED_REPAIRS_FOLDED_V4:
            metrics["exact_fixes"] += 1
            return _preserve_initial_case_v4(token, EXACT_GLUED_REPAIRS_FOLDED_V4[folded])
        if not _is_suspicious_glued_token_v4(token):
            return token
        connector_repaired, connector_changed = _split_connector_span_v4(token)
        if connector_changed:
            metrics["dynamic_segmentations"] += 1
            return connector_repaired
        repaired, changed, confidence = segment_glued_token_v4(token, SPANISH_WORDLIST_V4)
        if changed:
            metrics["dynamic_segmentations"] += 1
            return repaired
        if confidence > 0:
            metrics["low_confidence_skipped"] += 1
        return token

    for _ in range(3):
        previous = text
        text = re.sub(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ]+", repair_match, text)
        if text == previous:
            break
    after = detect_suspicious_glued_tokens(text)
    metrics["suspicious_after"] = after["suspicious_count"]
    if metrics["suspicious_before"]:
        metrics["suspicious_reduction_percent"] = (
            (metrics["suspicious_before"] - metrics["suspicious_after"])
            / metrics["suspicious_before"]
            * 100.0
        )
    return text, metrics


def repair_glued_words(text: str) -> str:
    """Backward-compatible wrapper around the v4 editorial repair."""
    repaired, _metrics = repair_glued_words_v4(text)
    return repaired

def fix_punctuation_spacing(text: str) -> str:
    """Add spaces after punctuation when missing."""
    text = re.sub(r'([,:;])([a-zA-ZáéíóúÁÉÍÓÚñÑ])', r'\1 \2', text)
    text = re.sub(r'(\.)([a-zA-ZáéíóúÁÉÍÓÚñÑ]{2,})', r'\1 \2', text)
    return text

def remove_repeated_running_headers(lines: list[str]) -> list[str]:
    """Detect and remove headers/footers that repeat across pages."""
    if len(lines) < 20:
        return lines
    counts = Counter(line.strip() for line in lines if line.strip() and len(line.strip()) < 50)
    running_headers = {line for line, count in counts.items() if count > 2}
    header_patterns = [
        r'Ars Magica', r'introducción', r'personajes', r'alianza',
        r'\d+$'
    ]
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        is_header = False
        if stripped in running_headers:
            for pattern in header_patterns:
                if re.search(pattern, stripped, re.I):
                    is_header = True
                    break
        if is_header:
            if line.startswith("#"):
                cleaned_lines.append(line)
            else:
                continue
        else:
            cleaned_lines.append(line)
    return cleaned_lines

def sanitize_markdown(content: str):
    """Orchestrate the editorial cleanup of Markdown."""
    content = remove_image_placeholders(content)
    content, _metrics = repair_glued_words_v4(content)
    content = fix_punctuation_spacing(content)
    
    lines = content.splitlines()
    lines = remove_repeated_running_headers(lines)
    
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append("")
            continue
        if stripped.startswith("![Image]") or stripped.startswith("![Table]"):
            continue
        if len(stripped) > 80 and re.match(r'^[a-zA-Z0-9+/=]+$', stripped):
            continue
        if len(stripped) == 1 and ord(stripped[0]) > 0x007F and not re.match(r'[¡¿áéíóúÁÉÍÓÚñÑüÜ]', stripped):
            continue
        line = line.replace("&amp;", "&").replace("&quot;", "\"").replace("&lt;", "<").replace("&gt;", ">")
        cleaned_lines.append(line)
        
    result = "\n".join(cleaned_lines)
    result = re.sub(r'\n{3,}', '\n\n', result)
    return result

def convert_md_to_docx(md_path: Path, docx_path: Path):
    if not md_path.exists():
        print(f"Error: Markdown file not found at {md_path}")
        sys.exit(1)
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        print("Error: python-docx not installed in this environment.")
        sys.exit(1)
    try:
        raw_content = md_path.read_text(encoding="utf-8")
    except Exception as e:
        print(f"Error reading Markdown: {e}")
        sys.exit(1)
    content = sanitize_markdown(raw_content)
    doc = Document()
    if 'Normal' in doc.styles:
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
    lines = content.splitlines()
    for line in lines:
        line_clean = line.strip()
        if not line_clean:
            continue
        if line_clean.startswith("# "):
            doc.add_heading(line_clean[2:], level=0)
        elif line_clean.startswith("## "):
            doc.add_heading(line_clean[3:], level=1)
        elif line_clean.startswith("### "):
            doc.add_heading(line_clean[4:], level=2)
        elif line_clean.startswith("#### "):
            doc.add_heading(line_clean[5:], level=3)
        elif line_clean.startswith("- ") or line_clean.startswith("* "):
            doc.add_paragraph(line_clean[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line_clean):
            text = re.sub(r'^\d+\. ', '', line_clean)
            doc.add_paragraph(text, style='List Number')
        elif line_clean.startswith("|") and "|" in line_clean[1:]:
            if re.match(r'^[ \|:\-]+$', line_clean):
                continue
            p = doc.add_paragraph(line_clean)
            try: p.style = 'Quote'
            except: pass
        else:
            doc.add_paragraph(line_clean)
    try:
        doc.save(str(docx_path))
        print(f"Successfully converted {md_path} to {docx_path}")
    except Exception as e:
        print(f"Error saving DOCX: {e}")
        sys.exit(1)

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("input", help="Input Markdown file")
    parser.add_argument("output", help="Output DOCX file")
    args = parser.parse_args()
    convert_md_to_docx(Path(args.input), Path(args.output))
